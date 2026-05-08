"""
pytest tests for the TenderPlatform pipeline.
Run from tender_platform/ directory:
    pytest tests/test_pipeline.py -v
"""

import sys
import os
import tempfile
from pathlib import Path
from datetime import datetime

# Ensure project root on path
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import pytest

# ──────────────────────────────────────────────
# Test data
# ──────────────────────────────────────────────

SAMPLE_TENDER_TEXT = """
Документация об аукционе. Федеральный закон № 44-ФЗ.
НМЦК: 3 500 000 рублей.
ИНН заказчика: 7712345678.

ТРЕБОВАНИЯ К УЧАСТНИКАМ:
1. Участник обязан иметь опыт исполнения аналогичных контрактов не менее 3 лет.
2. Необходимо наличие лицензии ФСТЭК России.
3. Участник должен подтвердить наличие сертификата соответствия ISO 27001.
4. Финансовое обеспечение заявки: 2% от НМЦК = 70 000 рублей.
5. Обеспечение исполнения контракта — 5% от НМЦК.
6. Среднесписочная численность — не менее 10 человек.
7. Выручка за последние 3 года — не менее 5 000 000 рублей ежегодно.
При нарушении сроков штраф 30% от стоимости контракта.
Национальный режим применяется в соответствии с Постановлением Правительства.
"""

SAMPLE_COMPANY = {
    "id": 1,
    "name": "ООО «Тест Компани»",
    "inn": "7712345678",
    "ogrn": "1187746123456",
    "legal_address": "г. Москва, ул. Тестовая, д. 1",
    "okpd2_codes": ["62.01", "62.02", "63.11"],
    "founding_year": 2016,
    "annual_revenue": 12_000_000,
    "annual_revenue_prev": 9_000_000,
    "staff_count": 45,
    "licenses": [
        {"name": "Лицензия ФСТЭК России на ТЗКИ", "type": "ФСТЭК"},
        {"name": "Сертификат ISO 27001", "type": "ISO"},
    ],
    "smp_status": True,
    "past_contracts": [
        {"customer": "Минцифры", "subject": "IT-услуги техническая поддержка", "amount": 3_500_000, "year": 2023},
        {"customer": "ФНС России", "subject": "Администрирование серверов", "amount": 2_800_000, "year": 2022},
        {"customer": "Росстат", "subject": "Техподдержка инфраструктуры", "amount": 1_900_000, "year": 2022},
    ],
    "authorized_person": "Иванов Иван Иванович",
    "authorized_person_position": "Генеральный директор",
    "bank_details": {"bank_name": "Сбербанк", "account": "40702810200000012345", "bik": "044525225"},
}

SAMPLE_TENDER = {
    "id": 1,
    "title": "Тестовый тендер IT-услуги",
    "external_id": "0372200097524000001",
    "customer_name": "Тестовый заказчик",
    "initial_price": 3_500_000,
    "law_type": "FZ_44",
    "purchase_subject": "Техническая поддержка ИТ-инфраструктуры",
    "full_text": SAMPLE_TENDER_TEXT,
    "okpd2_codes": ["62.02", "63.11"],
}


# ──────────────────────────────────────────────
# Tests
# ──────────────────────────────────────────────

class TestDocumentParser:
    def test_extracts_requirements_from_text(self):
        from modules.document_parser import DocumentParser

        parser = DocumentParser()
        parsed = {"full_text": SAMPLE_TENDER_TEXT, "pages": [{"number": 1, "text": SAMPLE_TENDER_TEXT}]}
        requirements = parser.extract_requirements(parsed)

        assert len(requirements) >= 3, (
            f"Expected >= 3 requirements, got {len(requirements)}: "
            f"{[r['text'][:60] for r in requirements]}"
        )

    def test_requirement_has_required_fields(self):
        from modules.document_parser import DocumentParser

        parser = DocumentParser()
        parsed = {"full_text": SAMPLE_TENDER_TEXT, "pages": [{"number": 1, "text": SAMPLE_TENDER_TEXT}]}
        requirements = parser.extract_requirements(parsed)

        for req in requirements:
            assert "category" in req
            assert "text" in req
            assert "is_mandatory" in req
            assert "risk_level" in req
            assert req["category"] in (
                "QUALIFICATION", "FINANCIAL", "TECHNICAL", "DOCUMENTATION", "RESTRICTION"
            )
            assert req["risk_level"] in ("LOW", "MEDIUM", "HIGH")

    def test_metadata_extraction(self):
        from modules.document_parser import DocumentParser

        parser = DocumentParser()
        metadata = parser.extract_metadata(SAMPLE_TENDER_TEXT)
        # Should extract at least law type
        assert "law_type" in metadata
        assert metadata["law_type"] == "FZ_44"

    def test_parse_txt_file(self):
        from modules.document_parser import DocumentParser

        parser = DocumentParser()
        with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", encoding="utf-8", delete=False) as f:
            f.write(SAMPLE_TENDER_TEXT)
            tmp_path = f.name
        try:
            result = parser.parse_file(tmp_path)
            assert "full_text" in result
            assert len(result["full_text"]) > 10
        finally:
            os.unlink(tmp_path)


class TestRequirementMatcher:
    def test_returns_score_between_0_and_1(self):
        from modules.document_parser import DocumentParser
        from modules.requirement_matcher import RequirementMatcher

        parser = DocumentParser()
        parsed = {"full_text": SAMPLE_TENDER_TEXT, "pages": [{"number": 1, "text": SAMPLE_TENDER_TEXT}]}
        requirements = parser.extract_requirements(parsed)

        matcher = RequirementMatcher()
        result = matcher.match(requirements, SAMPLE_COMPANY)

        assert 0.0 <= result["overall_match_score"] <= 1.0
        assert 0.0 <= result["qualification_score"] <= 1.0
        assert 0.0 <= result["financial_score"] <= 1.0
        assert 0.0 <= result["technical_score"] <= 1.0
        assert 0.0 <= result["documentation_score"] <= 1.0

    def test_returns_recommendation(self):
        from modules.document_parser import DocumentParser
        from modules.requirement_matcher import RequirementMatcher

        parser = DocumentParser()
        parsed = {"full_text": SAMPLE_TENDER_TEXT, "pages": [{"number": 1, "text": SAMPLE_TENDER_TEXT}]}
        requirements = parser.extract_requirements(parsed)

        matcher = RequirementMatcher()
        result = matcher.match(requirements, SAMPLE_COMPANY)

        assert result["recommendation"] in ("PARTICIPATE", "CONSIDER", "AVOID")

    def test_returns_matched_and_missing_lists(self):
        from modules.document_parser import DocumentParser
        from modules.requirement_matcher import RequirementMatcher

        parser = DocumentParser()
        parsed = {"full_text": SAMPLE_TENDER_TEXT, "pages": [{"number": 1, "text": SAMPLE_TENDER_TEXT}]}
        requirements = parser.extract_requirements(parsed)

        matcher = RequirementMatcher()
        result = matcher.match(requirements, SAMPLE_COMPANY)

        assert isinstance(result["matched_requirements"], list)
        assert isinstance(result["missing_requirements"], list)

    def test_empty_requirements_returns_valid_result(self):
        from modules.requirement_matcher import RequirementMatcher

        matcher = RequirementMatcher()
        result = matcher.match([], SAMPLE_COMPANY)
        assert result["overall_match_score"] == 1.0

    def test_company_capabilities_not_empty(self):
        from modules.requirement_matcher import RequirementMatcher

        matcher = RequirementMatcher()
        caps = matcher._company_capabilities_text(SAMPLE_COMPANY)
        assert len(caps) >= 3


class TestRiskAnalyzer:
    def test_detects_high_risk_penalty(self):
        from modules.risk_analyzer import RiskAnalyzer

        analyzer = RiskAnalyzer()
        text_with_penalty = (
            "При нарушении сроков поставки штраф составляет 30% от стоимости контракта. "
            "Участник обязан предоставить банковскую гарантию."
        )
        flags = analyzer.analyze(text_with_penalty, [])
        levels = [f["level"] for f in flags]
        assert "HIGH" in levels, f"Expected HIGH risk in: {flags}"

    def test_detects_national_regime(self):
        from modules.risk_analyzer import RiskAnalyzer

        analyzer = RiskAnalyzer()
        text = "Национальный режим применяется в соответствии с Постановлением Правительства РФ."
        flags = analyzer.analyze(text, [])
        descriptions = " ".join(f["description"] for f in flags)
        assert "национальный режим" in descriptions.lower() or any(
            "режим" in f["description"].lower() for f in flags
        )

    def test_overall_risk_level_high(self):
        from modules.risk_analyzer import RiskAnalyzer

        analyzer = RiskAnalyzer()
        flags = [
            {"level": "HIGH", "description": "Штраф 30%", "source_text": "...", "page": 1},
            {"level": "LOW", "description": "Гарантийный срок", "source_text": "...", "page": 1},
        ]
        assert analyzer.overall_risk_level(flags) == "HIGH"

    def test_overall_risk_level_medium(self):
        from modules.risk_analyzer import RiskAnalyzer

        analyzer = RiskAnalyzer()
        flags = [
            {"level": "MEDIUM", "description": "Неустойка", "source_text": "...", "page": 1},
            {"level": "LOW", "description": "Гарантия", "source_text": "...", "page": 1},
        ]
        assert analyzer.overall_risk_level(flags) == "MEDIUM"

    def test_overall_risk_level_low_on_empty(self):
        from modules.risk_analyzer import RiskAnalyzer

        analyzer = RiskAnalyzer()
        assert analyzer.overall_risk_level([]) == "LOW"

    def test_returns_list_of_dicts(self):
        from modules.risk_analyzer import RiskAnalyzer

        analyzer = RiskAnalyzer()
        flags = analyzer.analyze(SAMPLE_TENDER_TEXT, [])
        assert isinstance(flags, list)
        for f in flags:
            assert "level" in f
            assert "description" in f
            assert f["level"] in ("LOW", "MEDIUM", "HIGH")


class TestDocumentGenerator:
    def _make_analysis_dict(self):
        return {
            "id": 1,
            "overall_match_score": 0.75,
            "qualification_score": 0.80,
            "financial_score": 0.70,
            "technical_score": 0.75,
            "documentation_score": 0.90,
            "win_probability": 0.55,
            "matched_requirements": [
                {"text": "Лицензия ФСТЭК", "category": "QUALIFICATION", "is_mandatory": True, "similarity": 0.8},
            ],
            "missing_requirements": [],
            "risk_flags": [],
            "recommendation": "PARTICIPATE",
        }

    def test_creates_application_docx(self):
        from modules.doc_generator import DocumentGenerator

        gen = DocumentGenerator()
        path = gen.generate_application(SAMPLE_TENDER, SAMPLE_COMPANY, self._make_analysis_dict())
        assert Path(path).exists(), f"File not found: {path}"
        assert path.endswith(".docx")

    def test_creates_declaration_docx(self):
        from modules.doc_generator import DocumentGenerator

        gen = DocumentGenerator()
        path = gen.generate_declaration(SAMPLE_TENDER, SAMPLE_COMPANY)
        assert Path(path).exists()
        assert path.endswith(".docx")

    def test_creates_technical_proposal_docx(self):
        from modules.doc_generator import DocumentGenerator

        gen = DocumentGenerator()
        path = gen.generate_technical_proposal(SAMPLE_TENDER, SAMPLE_COMPANY, self._make_analysis_dict())
        assert Path(path).exists()
        assert path.endswith(".docx")

    def test_creates_form2_docx(self):
        from modules.doc_generator import DocumentGenerator

        gen = DocumentGenerator()
        path = gen.generate_form_2(SAMPLE_TENDER, SAMPLE_COMPANY)
        assert Path(path).exists()
        assert path.endswith(".docx")

    def test_docx_is_valid_word_file(self):
        from modules.doc_generator import DocumentGenerator
        from docx import Document

        gen = DocumentGenerator()
        path = gen.generate_application(SAMPLE_TENDER, SAMPLE_COMPANY, self._make_analysis_dict())
        doc = Document(path)
        # Valid DOCX should have at least one paragraph
        assert len(doc.paragraphs) >= 1

    def test_full_package_returns_four_files(self):
        from modules.doc_generator import DocumentGenerator

        gen = DocumentGenerator()
        paths = gen.generate_full_package(SAMPLE_TENDER, SAMPLE_COMPANY, self._make_analysis_dict())
        assert len(paths) == 4
        for p in paths:
            assert Path(p).exists()


class TestWinProbabilityModel:
    def test_predict_returns_float_0_1(self):
        from modules.win_probability_model import WinProbabilityModel

        model = WinProbabilityModel()
        analysis = {
            "overall_match_score": 0.75,
            "qualification_score": 0.80,
            "financial_score": 0.70,
            "technical_score": 0.75,
            "documentation_score": 0.90,
            "matched_requirements": [],
            "missing_requirements": [],
            "risk_flags": [],
        }
        prob = model.predict(analysis, SAMPLE_TENDER, SAMPLE_COMPANY)
        assert 0.0 <= prob <= 1.0, f"Probability {prob} out of range"

    def test_higher_match_gives_higher_probability(self):
        from modules.win_probability_model import WinProbabilityModel

        model = WinProbabilityModel()
        good_analysis = {
            "overall_match_score": 0.95,
            "qualification_score": 0.95,
            "financial_score": 0.95,
            "technical_score": 0.95,
            "documentation_score": 0.95,
            "matched_requirements": [],
            "missing_requirements": [],
            "risk_flags": [],
        }
        bad_analysis = {
            "overall_match_score": 0.1,
            "qualification_score": 0.1,
            "financial_score": 0.1,
            "technical_score": 0.1,
            "documentation_score": 0.1,
            "matched_requirements": [],
            "missing_requirements": [
                {"is_mandatory": True},
                {"is_mandatory": True},
                {"is_mandatory": True},
            ],
            "risk_flags": [{"level": "HIGH"}, {"level": "HIGH"}],
        }
        good_prob = model.predict(good_analysis, SAMPLE_TENDER, SAMPLE_COMPANY)
        bad_prob = model.predict(bad_analysis, SAMPLE_TENDER, SAMPLE_COMPANY)
        assert good_prob > bad_prob, f"Expected {good_prob:.3f} > {bad_prob:.3f}"

    def test_get_features_correct_shape(self):
        from modules.win_probability_model import WinProbabilityModel

        model = WinProbabilityModel()
        analysis = {
            "overall_match_score": 0.7,
            "qualification_score": 0.7,
            "financial_score": 0.7,
            "technical_score": 0.7,
            "documentation_score": 0.7,
            "matched_requirements": [],
            "missing_requirements": [],
            "risk_flags": [],
        }
        features = model.get_features(analysis, SAMPLE_TENDER, SAMPLE_COMPANY)
        assert features.shape == (13,)

    def test_train_on_synthetic_data(self):
        from modules.win_probability_model import WinProbabilityModel

        model = WinProbabilityModel()
        result = model.train([])  # triggers synthetic data training
        assert result["trained"] is True
        assert result["n_samples"] > 0


class TestFullPipeline:
    def test_full_pipeline_from_text_to_document(self):
        """End-to-end: parse text → match → risk → win_prob → generate document."""
        from modules.document_parser import DocumentParser
        from modules.requirement_matcher import RequirementMatcher
        from modules.risk_analyzer import RiskAnalyzer
        from modules.win_probability_model import WinProbabilityModel
        from modules.doc_generator import DocumentGenerator

        # Step 1: parse
        parser = DocumentParser()
        parsed = {
            "full_text": SAMPLE_TENDER_TEXT,
            "pages": [{"number": 1, "text": SAMPLE_TENDER_TEXT}],
        }
        requirements = parser.extract_requirements(parsed)
        assert len(requirements) >= 1

        # Step 2: match
        matcher = RequirementMatcher()
        match_result = matcher.match(requirements, SAMPLE_COMPANY)
        assert 0.0 <= match_result["overall_match_score"] <= 1.0

        # Step 3: risk
        risk_analyzer = RiskAnalyzer()
        risk_flags = risk_analyzer.analyze(SAMPLE_TENDER_TEXT, requirements)
        match_result["risk_flags"] = risk_flags

        # Step 4: win probability
        win_model = WinProbabilityModel()
        win_prob = win_model.predict(match_result, SAMPLE_TENDER, SAMPLE_COMPANY)
        assert 0.0 <= win_prob <= 1.0

        # Step 5: generate document
        analysis_dict = {
            **match_result,
            "id": 99,
            "win_probability": win_prob,
        }
        gen = DocumentGenerator()
        path = gen.generate_application(SAMPLE_TENDER, SAMPLE_COMPANY, analysis_dict)
        assert Path(path).exists()

        # Verify it's a valid DOCX
        from docx import Document
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text)
        assert "ЗАЯВКА" in text.upper() or len(text) > 20

    def test_pipeline_with_database(self, tmp_path):
        """Pipeline with SQLite in-memory database."""
        import os
        os.environ["DATABASE_URL"] = f"sqlite:///{tmp_path}/test.db"

        from sqlalchemy import create_engine
        from sqlalchemy.orm import sessionmaker
        from core.models import Base, Tender, TenderRequirement, CompanyProfile, TenderAnalysis
        from modules.document_parser import DocumentParser
        from modules.requirement_matcher import RequirementMatcher
        from modules.risk_analyzer import RiskAnalyzer
        from modules.win_probability_model import WinProbabilityModel

        engine = create_engine(f"sqlite:///{tmp_path}/test_pipeline.db", echo=False)
        Base.metadata.create_all(engine)
        Session = sessionmaker(bind=engine)
        db = Session()

        try:
            # Create company
            company = CompanyProfile(
                name="Тест ООО",
                inn="7712345678",
                okpd2_codes=["62.01"],
                founding_year=2015,
                annual_revenue=10_000_000,
                staff_count=20,
                licenses=[{"name": "Лицензия ФСТЭК"}],
                smp_status=True,
                past_contracts=[{"customer": "МВД", "subject": "IT-услуги", "amount": 2000000, "year": 2023}],
            )
            db.add(company)

            # Create tender
            tender = Tender(
                title="Тестовый тендер",
                full_text=SAMPLE_TENDER_TEXT,
                law_type="FZ_44",
                initial_price=3_500_000,
                status="NEW",
            )
            db.add(tender)
            db.flush()

            # Run pipeline
            parser = DocumentParser()
            parsed_doc = {"full_text": tender.full_text, "pages": [{"number": 1, "text": tender.full_text}]}
            req_dicts = parser.extract_requirements(parsed_doc)

            for rd in req_dicts:
                db.add(TenderRequirement(
                    tender_id=tender.id,
                    category=rd["category"],
                    text=rd["text"],
                    is_mandatory=rd["is_mandatory"],
                    source_page=rd.get("source_page"),
                    extracted_value=rd.get("extracted_value"),
                    risk_level=rd.get("risk_level", "LOW"),
                ))
            db.flush()

            company_dict = {
                "id": company.id, "name": company.name, "inn": company.inn,
                "okpd2_codes": company.okpd2_codes, "founding_year": company.founding_year,
                "annual_revenue": float(company.annual_revenue), "staff_count": company.staff_count,
                "licenses": company.licenses, "smp_status": company.smp_status,
                "past_contracts": company.past_contracts, "ogrn": None,
                "legal_address": None, "annual_revenue_prev": None,
                "authorized_person": None, "authorized_person_position": None, "bank_details": None,
            }

            matcher = RequirementMatcher()
            match_result = matcher.match(req_dicts, company_dict)

            risk_analyzer = RiskAnalyzer()
            risk_flags = risk_analyzer.analyze(tender.full_text, req_dicts)
            match_result["risk_flags"] = risk_flags

            win_model = WinProbabilityModel()
            tender_dict = {
                "id": tender.id, "title": tender.title, "initial_price": float(tender.initial_price),
                "law_type": tender.law_type, "full_text": tender.full_text, "okpd2_codes": None,
            }
            win_prob = win_model.predict(match_result, tender_dict, company_dict)

            analysis = TenderAnalysis(
                tender_id=tender.id,
                company_id=company.id,
                overall_match_score=match_result["overall_match_score"],
                win_probability=win_prob,
                recommendation=match_result["recommendation"],
                matched_requirements=match_result["matched_requirements"],
                missing_requirements=match_result["missing_requirements"],
                risk_flags=risk_flags,
            )
            db.add(analysis)
            db.commit()

            saved = db.query(TenderAnalysis).filter_by(tender_id=tender.id).first()
            assert saved is not None
            assert saved.overall_match_score is not None
            assert saved.win_probability is not None
            assert saved.recommendation in ("PARTICIPATE", "CONSIDER", "AVOID")

        finally:
            db.close()
