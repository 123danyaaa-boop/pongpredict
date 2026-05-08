import logging
import re
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional

logger = logging.getLogger(__name__)

TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates"
OUTPUT_DIR = Path(__file__).resolve().parent.parent / "output" / "generated_docs"


def _ensure_output_dir():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def _fill_placeholders(text: str, context: Dict[str, Any]) -> str:
    def replacer(m):
        key = m.group(1)
        return str(context.get(key, m.group(0)))

    return re.sub(r"\{\{(\w+)\}\}", replacer, text)


def _add_heading(doc, text: str, level: int = 1):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT


def _add_paragraph(doc, text: str, bold: bool = False, indent: bool = False):
    from docx.shared import Pt, Cm
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.bold = bold
    return p


def _add_table_row(table, cells: List[str]):
    row = table.add_row()
    for i, cell_text in enumerate(cells):
        row.cells[i].text = cell_text
    return row


class DocumentGenerator:
    def generate_full_package(
        self,
        tender: Dict[str, Any],
        company: Dict[str, Any],
        analysis: Dict[str, Any],
    ) -> List[str]:
        _ensure_output_dir()
        paths = []
        paths.append(self.generate_application(tender, company, analysis))
        paths.append(self.generate_declaration(tender, company))
        paths.append(self.generate_technical_proposal(tender, company, analysis))
        paths.append(self.generate_form_2(tender, company))
        return paths

    def generate_application(
        self,
        tender: Dict[str, Any],
        company: Dict[str, Any],
        analysis: Dict[str, Any],
    ) -> str:
        try:
            from docx import Document
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx is required: pip install python-docx")

        doc = Document()
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(12)

        _add_heading(doc, "ЗАЯВКА НА УЧАСТИЕ В ЗАКУПКЕ", level=1)
        _add_heading(
            doc,
            f"{tender.get('title', 'Наименование закупки')}",
            level=2,
        )

        doc.add_paragraph()
        _add_paragraph(
            doc,
            f"Дата составления: {datetime.now().strftime('%d.%m.%Y')}",
        )
        doc.add_paragraph()

        _add_heading(doc, "1. Сведения об участнике закупки", level=2)
        info_table = doc.add_table(rows=1, cols=2)
        info_table.style = "Table Grid"
        _add_table_row(info_table, ["Наименование организации", company.get("name", "")])
        _add_table_row(info_table, ["ИНН", company.get("inn", "")])
        _add_table_row(info_table, ["ОГРН", company.get("ogrn", "")])
        _add_table_row(info_table, ["Юридический адрес", company.get("legal_address", "")])
        bd = company.get("bank_details") or {}
        if isinstance(bd, dict):
            _add_table_row(info_table, ["Банк", bd.get("bank_name", "")])
            _add_table_row(info_table, ["Расчётный счёт", bd.get("account", "")])
            _add_table_row(info_table, ["БИК", bd.get("bik", "")])

        doc.add_paragraph()
        _add_heading(doc, "2. Согласие с условиями закупки", level=2)
        law_type = tender.get("law_type", "FZ_44")
        law_display = "44-ФЗ" if law_type == "FZ_44" else "223-ФЗ"
        _add_paragraph(
            doc,
            f"Участник закупки ({company.get('name', '')}) подтверждает своё безоговорочное согласие "
            f"с условиями проведения закупки в соответствии с {law_display} и документацией о закупке.",
        )

        doc.add_paragraph()
        _add_heading(doc, "3. Соответствие требованиям заказчика", level=2)
        req_table = doc.add_table(rows=1, cols=3)
        req_table.style = "Table Grid"
        header = req_table.rows[0]
        header.cells[0].text = "Требование заказчика"
        header.cells[1].text = "Подтверждение компании"
        header.cells[2].text = "Подтверждающий документ"

        matched = analysis.get("matched_requirements") or []
        for req in matched[:15]:
            row = req_table.add_row()
            row.cells[0].text = req.get("text", "")[:150]
            row.cells[1].text = "Соответствует"
            row.cells[2].text = "Приложение к заявке"

        doc.add_paragraph()
        _add_heading(doc, "4. Подпись уполномоченного лица", level=2)
        sign_table = doc.add_table(rows=1, cols=2)
        sign_table.style = "Table Grid"
        _add_table_row(sign_table, ["ФИО", company.get("authorized_person", "")])
        _add_table_row(sign_table, ["Должность", company.get("authorized_person_position", "")])
        _add_table_row(sign_table, ["Дата", datetime.now().strftime("%d.%m.%Y")])
        _add_table_row(sign_table, ["Подпись", "___________________________"])

        filename = f"application_{tender.get('id', 'new')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        out_path = OUTPUT_DIR / filename
        doc.save(str(out_path))
        logger.info(f"Application saved: {out_path}")
        return str(out_path)

    def generate_declaration(
        self,
        tender: Dict[str, Any],
        company: Dict[str, Any],
    ) -> str:
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            raise ImportError("python-docx is required: pip install python-docx")

        doc = Document()
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(12)

        law_type = tender.get("law_type", "FZ_44")
        law_display = "44-ФЗ" if law_type == "FZ_44" else "223-ФЗ"

        _add_heading(doc, f"ДЕКЛАРАЦИЯ СООТВЕТСТВИЯ ТРЕБОВАНИЯМ {law_display}", level=1)
        doc.add_paragraph()

        _add_paragraph(
            doc,
            f"Я, {company.get('authorized_person', '______________')}, "
            f"{company.get('authorized_person_position', '______________')} "
            f"организации {company.get('name', '______________')} "
            f"(ИНН: {company.get('inn', '______________')}), "
            f"настоящим подтверждаю соответствие участника закупки следующим требованиям:",
        )
        doc.add_paragraph()

        DECLARATION_ITEMS_44 = [
            "Соответствие требованиям, установленным в соответствии с законодательством Российской Федерации к лицам, осуществляющим поставку товара, выполнение работы, оказание услуги, являющихся объектом закупки",
            "Непроведение ликвидации участника закупки — юридического лица и отсутствие решения арбитражного суда о признании участника закупки банкротом",
            "Неприостановление деятельности участника закупки в порядке, установленном КоАП РФ",
            "Отсутствие у участника закупки недоимки по налогам, сборам, задолженности по иным обязательным платежам",
            "Отсутствие у руководителя, членов коллегиального исполнительного органа участника закупки судимости за преступления в сфере экономики",
            "Обладание участником закупки исключительными правами на результаты интеллектуальной деятельности (если требуется)",
            "Отсутствие между участником закупки и заказчиком конфликта интересов",
            "Участник закупки не является офшорной компанией",
        ]

        items = DECLARATION_ITEMS_44 if law_type == "FZ_44" else DECLARATION_ITEMS_44
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"✓ {item}")

        doc.add_paragraph()
        _add_paragraph(
            doc,
            f"Дата: {datetime.now().strftime('%d.%m.%Y')}",
        )
        _add_paragraph(
            doc,
            f"Подпись: _________________________ / {company.get('authorized_person', '')} /",
        )
        _add_paragraph(doc, "М.П.")

        filename = f"declaration_{tender.get('id', 'new')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        out_path = OUTPUT_DIR / filename
        doc.save(str(out_path))
        logger.info(f"Declaration saved: {out_path}")
        return str(out_path)

    def generate_technical_proposal(
        self,
        tender: Dict[str, Any],
        company: Dict[str, Any],
        analysis: Dict[str, Any],
    ) -> str:
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            raise ImportError("python-docx is required: pip install python-docx")

        doc = Document()
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(12)

        _add_heading(doc, "ТЕХНИЧЕСКОЕ ПРЕДЛОЖЕНИЕ", level=1)
        _add_heading(
            doc,
            f"На участие в закупке: {tender.get('title', '')}",
            level=2,
        )
        doc.add_paragraph()

        _add_heading(doc, "1. Описание предлагаемого товара/услуги/работы", level=2)
        subject = tender.get("purchase_subject") or tender.get("title", "")
        _add_paragraph(
            doc,
            f"Организация {company.get('name', '')} предлагает к поставке/выполнению: {subject}. "
            f"Предложение полностью соответствует техническому заданию заказчика.",
        )

        doc.add_paragraph()
        _add_heading(doc, "2. Соответствие техническим характеристикам", level=2)
        tech_reqs = [
            r for r in (analysis.get("matched_requirements") or [])
            if r.get("category") in ("TECHNICAL", "QUALIFICATION")
        ]
        if tech_reqs:
            tech_table = doc.add_table(rows=1, cols=3)
            tech_table.style = "Table Grid"
            header = tech_table.rows[0]
            header.cells[0].text = "Требование ТЗ"
            header.cells[1].text = "Предложение участника"
            header.cells[2].text = "Соответствует"
            for req in tech_reqs[:20]:
                row = tech_table.add_row()
                row.cells[0].text = req.get("text", "")[:150]
                row.cells[1].text = "Соответствует требованиям"
                row.cells[2].text = "Да"
        else:
            _add_paragraph(doc, "Все технические требования заказчика выполняются в полном объёме.")

        doc.add_paragraph()
        _add_heading(doc, "3. Опыт исполнения аналогичных контрактов", level=2)
        past_contracts = company.get("past_contracts") or []
        if past_contracts:
            contracts_table = doc.add_table(rows=1, cols=4)
            contracts_table.style = "Table Grid"
            header = contracts_table.rows[0]
            header.cells[0].text = "Заказчик"
            header.cells[1].text = "Предмет контракта"
            header.cells[2].text = "Сумма, руб."
            header.cells[3].text = "Год"
            for contract in past_contracts[:10]:
                if isinstance(contract, dict):
                    row = contracts_table.add_row()
                    row.cells[0].text = str(contract.get("customer", ""))
                    row.cells[1].text = str(contract.get("subject", ""))
                    row.cells[2].text = str(contract.get("amount", ""))
                    row.cells[3].text = str(contract.get("year", ""))
        else:
            _add_paragraph(doc, "Информация о прошлых контрактах не предоставлена.")

        doc.add_paragraph()
        _add_paragraph(
            doc,
            f"Уполномоченное лицо: {company.get('authorized_person', '')} "
            f"({company.get('authorized_person_position', '')})",
        )
        _add_paragraph(doc, f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
        _add_paragraph(doc, "Подпись: _________________________")

        filename = f"technical_proposal_{tender.get('id', 'new')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        out_path = OUTPUT_DIR / filename
        doc.save(str(out_path))
        logger.info(f"Technical proposal saved: {out_path}")
        return str(out_path)

    def generate_form_2(
        self,
        tender: Dict[str, Any],
        company: Dict[str, Any],
    ) -> str:
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            raise ImportError("python-docx is required: pip install python-docx")

        doc = Document()
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(12)

        _add_heading(doc, 'ФОРМА 2. СВЕДЕНИЯ ОБ УЧАСТНИКЕ ЗАКУПКИ', level=1)
        doc.add_paragraph()

        info_table = doc.add_table(rows=1, cols=2)
        info_table.style = "Table Grid"
        rows_data = [
            ("Полное наименование", company.get("name", "")),
            ("ИНН", company.get("inn", "")),
            ("ОГРН", company.get("ogrn", "")),
            ("Юридический адрес", company.get("legal_address", "")),
            ("Год основания", str(company.get("founding_year", ""))),
            ("Среднесписочная численность", str(company.get("staff_count", ""))),
            ("Годовая выручка, руб.", str(company.get("annual_revenue", ""))),
            (
                "Субъект МСП",
                "Да" if company.get("smp_status") else "Нет",
            ),
            ("Уполномоченное лицо", company.get("authorized_person", "")),
            ("Должность", company.get("authorized_person_position", "")),
        ]
        for label, value in rows_data:
            _add_table_row(info_table, [label, value])

        bd = company.get("bank_details") or {}
        if isinstance(bd, dict):
            _add_table_row(info_table, ["Наименование банка", bd.get("bank_name", "")])
            _add_table_row(info_table, ["Расчётный счёт", bd.get("account", "")])
            _add_table_row(info_table, ["Корр. счёт", bd.get("corr_account", "")])
            _add_table_row(info_table, ["БИК", bd.get("bik", "")])

        doc.add_paragraph()
        _add_paragraph(doc, f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
        _add_paragraph(doc, "Подпись: _________________________  М.П.")

        filename = f"form2_{tender.get('id', 'new')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        out_path = OUTPUT_DIR / filename
        doc.save(str(out_path))
        logger.info(f"Form 2 saved: {out_path}")
        return str(out_path)

    def _fill_template(self, template_path: str, context: Dict[str, Any]):
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required: pip install python-docx")

        doc = Document(template_path)
        for para in doc.paragraphs:
            for run in para.runs:
                run.text = _fill_placeholders(run.text, context)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = _fill_placeholders(run.text, context)
        return doc
